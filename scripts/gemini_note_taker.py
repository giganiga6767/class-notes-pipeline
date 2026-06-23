#!/usr/bin/env python3
"""
gemini_note_taker.py
Cleans up raw speech-to-text dumps into structured markdown notes.
"""

import os
import sys
import re
import time
from pathlib import Path
from google import genai
from google.genai.errors import APIError

# ── Config ──────────────────────────────────────────────────────────────────
MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
API_KEY = os.environ.get("GEMINI_API_KEY")
if not API_KEY:
    try:
        client = genai.Client()
    except Exception:
        sys.exit("❌  Please set GEMINI_API_KEY environment variable.")
else:
    client = genai.Client(api_key=API_KEY)

def generate_content_with_retry(client, model, contents, max_retries=5, backoff=2.0):
    delay = 2.0
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(model=model, contents=contents)
        except APIError as e:
            code = getattr(e, "code", None)
            is_transient = code in (429, 503) or any(err_str in str(e) for err_str in ("503", "429", "UNAVAILABLE", "Resource Exhausted", "RESOURCE_EXHAUSTED"))
            if is_transient and attempt < max_retries - 1:
                print(f"⚠️  Gemini API temporary error ({code or '503/429'}). Retrying in {delay:.1f} seconds (attempt {attempt + 1}/{max_retries})...")
                time.sleep(delay)
                delay *= backoff
            else:
                raise

# ── Prompt ───────────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """
You are a professional note-taker and editor. You receive raw, messy text from
speech-to-text engines — full of filler words, repetition, broken sentences,
and no punctuation. Your job is to produce clean, beautifully structured
markdown notes.

Rules:
1. Remove ALL filler: "um", "uh", "like", "you know", "so yeah", "basically", etc.
2. Fix grammar, punctuation, and capitalization.
3. Identify and group distinct TOPICS. Give each a clear ## Heading.
4. Extract any action items or decisions into a separate ## Action Items section.
5. Write a 2-3 sentence ## Summary at the top.
6. Preserve the speaker's intent — do not add your own opinions.
7. If dates, names, or numbers appear, format them cleanly.
8. Output ONLY valid markdown. No preamble, no explanation.
"""

def clean_notes(raw_text: str, context: str = "") -> str:
    user_msg = f"""Please clean up and structure these raw notes.
{f'Context: {context}' if context else ''}

RAW TEXT:
---
{raw_text}
---
"""
    response = generate_content_with_retry(
        client=client,
        model=MODEL,
        contents=[{"role": "user", "parts": [{"text": SYSTEM_PROMPT + "\n\n" + user_msg}]}]
    )
    return response.text

def chunk_text(text: str, max_chars: int = 12_000) -> list[str]:
    """Split on paragraph boundaries, not mid-sentence."""
    paras = text.split("\n\n")
    chunks, current = [], ""
    for p in paras:
        if len(current) + len(p) > max_chars:
            chunks.append(current.strip())
            current = p
        else:
            current += "\n\n" + p
    if current:
        chunks.append(current.strip())
    return chunks

def process_notes(raw_text: str, context: str = "") -> str:
    if len(raw_text) <= 15000:
        return clean_notes(raw_text, context)
    
    print(f"  📝 Content length ({len(raw_text):,} chars) exceeds single prompt limit. Chunking...")
    chunks = chunk_text(raw_text)
    print(f"  🧩 Split into {len(chunks)} chunks.")
    parts = []
    for idx, chunk in enumerate(chunks):
        print(f"  ⚡ Processing chunk {idx + 1}/{len(chunks)}...")
        parts.append(clean_notes(chunk, context))
        
    print("  🤝 Merging all processed chunks...")
    merged = "\n\n---\n\n".join(parts)
    merge_prompt = f"""Merge these note sections into one unified document, removing duplicates, aligning headers, and consolidating action items:
    
{merged}"""
    
    response = generate_content_with_retry(
        client=client,
        model=MODEL,
        contents=[{"role": "user", "parts": [{"text": SYSTEM_PROMPT + "\n\n" + merge_prompt}]}]
    )
    return response.text

def parse_inline_formatting(paragraph, text, set_run_font):
    import re
    # Regex to split on bold/italic markup
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            bold_text = token[2:-2]
            run = paragraph.add_run(bold_text)
            set_run_font(run, bold=True)
        elif token.startswith('*') and token.endswith('*'):
            italic_text = token[1:-1]
            run = paragraph.add_run(italic_text)
            set_run_font(run, italic=True)
        else:
            run = paragraph.add_run(token)
            set_run_font(run)

def save_notes_as_docx(markdown_content: str, docx_path: Path):
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    PRIMARY_COLOR = (49, 46, 129)    # Deep Indigo
    TEXT_COLOR = (31, 41, 55)       # Dark Gray
    ACCENT_COLOR = (79, 70, 229)    # Indigo Accent
    
    def set_run_font(run, font_name="Arial", size_pt=11, color_rgb=TEXT_COLOR, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(*color_rgb)
        run.bold = bold
        run.italic = italic
        
    # Clean Markdown Parsing line-by-line
    lines = markdown_content.split('\n')
    
    # Find Document Title
    title_text = "Class Notes"
    for line in lines:
        if line.strip().startswith('# '):
            title_text = line.strip('# ').strip()
            break
            
    # Add Document Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(title_text)
    set_run_font(title_run, size_pt=22, color_rgb=PRIMARY_COLOR, bold=True)
    
    # Subtitle with date
    from datetime import datetime
    date_str = datetime.now().strftime("%B %d, %Y")
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run(f"Formatted on {date_str} • Powered by Gemini AI")
    set_run_font(sub_run, size_pt=10, color_rgb=(100, 116, 139), italic=True)
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # Skip the title line since we already added it
        if stripped.startswith('# ') and stripped.strip('# ').strip() == title_text:
            continue
            
        # Headings
        if stripped.startswith('# '):
            heading_text = stripped.strip('# ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            set_run_font(run, size_pt=18, color_rgb=PRIMARY_COLOR, bold=True)
            
        elif stripped.startswith('## '):
            heading_text = stripped.strip('# ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            set_run_font(run, size_pt=14, color_rgb=ACCENT_COLOR, bold=True)
            
        elif stripped.startswith('### '):
            heading_text = stripped.strip('# ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            set_run_font(run, size_pt=12, color_rgb=TEXT_COLOR, bold=True)
            
        # Bullet Lists
        elif stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            is_ordered = bool(re.match(r'^\d+\.\s', stripped))
            if is_ordered:
                prefix_match = re.match(r'^(\d+\.\s)', stripped)
                bullet_content = stripped[len(prefix_match.group(1)):].strip()
                list_style = 'List Number'
            else:
                bullet_content = stripped[2:].strip()
                list_style = 'List Bullet'
                
            p = doc.add_paragraph(style=list_style)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            parse_inline_formatting(p, bullet_content, set_run_font)
            
        # Normal Paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, stripped, set_run_font)
            
    doc.save(str(docx_path))

# ── Main ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python gemini_note_taker.py <input_file.txt> [context]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    context    = sys.argv[2] if len(sys.argv) > 2 else ""
    output_path = input_path.with_suffix(".notes.md")

    if not input_path.exists():
        sys.exit(f"❌ Input file not found: {input_path}")

    raw = input_path.read_text(encoding="utf-8")
    print(f"📄 Processing {len(raw):,} characters from {input_path.name}...")

    result = process_notes(raw, context)

    output_path.write_text(result, encoding="utf-8")
    print(f"✅  Clean notes saved → {output_path}")
    
    docx_path = input_path.with_suffix(".notes.docx")
    try:
        save_notes_as_docx(result, docx_path)
        print(f"✅  Word document notes saved → {docx_path}")
    except Exception as e:
        print(f"⚠️  Failed to save DOCX notes: {e}")
        
    print("\n" + "─"*60)
    print(result[:800] + ("..." if len(result) > 800 else ""))
