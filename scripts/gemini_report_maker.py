#!/usr/bin/env python3
"""
gemini_report_maker.py
Beautiful report from mixed files using Gemini, python-docx, and Matplotlib.
Generates both DOCX and PDF (via LibreOffice).
"""

import os, sys, base64, json, re, tempfile, time
from pathlib import Path
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from google import genai
from google.genai.errors import APIError
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ── Config ───────────────────────────────────────────────────────────────────
MODEL   = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
API_KEY = os.environ.get("GEMINI_API_KEY")
if not API_KEY:
    try:
        client = genai.Client()
    except Exception:
        sys.exit("❌  Please set GEMINI_API_KEY environment variable.")
else:
    client = genai.Client(api_key=API_KEY)

def generate_content_with_retry(client, model, contents, max_retries=5, backoff=2.0):
    delay = 2.0
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(model=model, contents=contents)
        except APIError as e:
            code = getattr(e, "code", None)
            is_transient = code in (429, 503) or any(err_str in str(e) for err_str in ("503", "429", "UNAVAILABLE", "Resource Exhausted", "RESOURCE_EXHAUSTED"))
            if is_transient and attempt < max_retries - 1:
                print(f"⚠️  Gemini API temporary error ({code or '503/429'}). Retrying in {delay:.1f} seconds (attempt {attempt + 1}/{max_retries})...")
                time.sleep(delay)
                delay *= backoff
            else:
                raise

# ── 1. File Ingestion ─────────────────────────────────────────────────────────
def ingest_files(file_paths: list[str]) -> dict:
    """Read all provided files into a structured context dict."""
    context = {"texts": [], "dataframes": {}, "images": [], "raw_data": []}

    for fp in file_paths:
        p = Path(fp)
        if not p.exists():
            print(f"⚠️  File not found: {fp}")
            continue
        ext = p.suffix.lower()

        if ext in {".txt", ".md", ".log"}:
            text = p.read_text(encoding="utf-8", errors="replace")
            context["texts"].append({"name": p.name, "content": text})
            print(f"  📄 Text file loaded: {p.name} ({len(text):,} chars)")

        elif ext == ".csv":
            df = pd.read_csv(p)
            context["dataframes"][p.name] = df
            context["raw_data"].append({"name": p.name, "preview": df.to_string(max_rows=20)})
            print(f"  📊 CSV loaded: {p.name} ({len(df)} rows × {len(df.columns)} cols)")

        elif ext == ".json":
            data = json.loads(p.read_text())
            if isinstance(data, list):
                df = pd.DataFrame(data)
                context["dataframes"][p.name] = df
                context["raw_data"].append({"name": p.name, "preview": df.to_string(max_rows=20)})
            else:
                context["texts"].append({"name": p.name, "content": json.dumps(data, indent=2)})
            print(f"  🗂️  JSON loaded: {p.name}")

        elif ext in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
            b64 = base64.b64encode(p.read_bytes()).decode()
            mime = {"png":"image/png","jpg":"image/jpeg","jpeg":"image/jpeg",
                    "gif":"image/gif","webp":"image/webp"}[ext.lstrip(".")]
            context["images"].append({"name": p.name, "b64": b64, "mime": mime})
            print(f"  🖼️  Image loaded: {p.name}")

    return context


# ── 2. Gemini Analysis ────────────────────────────────────────────────────────
ANALYSIS_PROMPT = """
You are a professional data analyst and report writer. You will receive:
- Text content from various files
- Previews of data tables (CSV/JSON)
- Descriptions of images provided

Your job is to return a JSON object (and ONLY a JSON object, no markdown fences) with:
{
  "report_title": "A compelling title for the overall report",
  "executive_summary": "3-5 sentence executive summary of all the content",
  "sections": [
    {
      "title": "Section Title",
      "content": "Well-written prose for this section (2-4 paragraphs)",
      "key_insights": ["insight 1", "insight 2"]
    }
  ],
  "chart_specs": [
    {
      "chart_type": "pie|bar|line|scatter|heatmap|histogram",
      "title": "Chart title",
      "description": "What this chart shows",
      "data_source": "which file/data this is based on",
      "x_col": "column name for x axis (bar/line/scatter)",
      "y_col": "column name for y axis (bar/line/scatter)",
      "label_col": "column name for labels (pie)",
      "value_col": "column name for values (pie)",
      "color_col": "optional grouping column"
    }
  ],
  "conclusion": "Strong 2-3 paragraph conclusion and recommendations",
  "image_captions": {"filename.png": "Caption for this image"}
}

The chart_specs must only reference columns that actually exist in the data provided.
"""

def analyze_with_gemini(context: dict, report_topic: str = "") -> dict:
    """Send all content to Gemini for analysis and structure."""
    parts = []

    # Build text summary of all content
    summary_parts = []
    if report_topic:
        summary_parts.append(f"Report Topic / User Request: {report_topic}\n")

    for t in context["texts"]:
        summary_parts.append(f"=== FILE: {t['name']} ===\n{t['content'][:3000]}\n")

    for rd in context["raw_data"]:
        summary_parts.append(f"=== DATA FILE: {rd['name']} ===\n{rd['preview']}\n")

    combined_text = "\n".join(summary_parts)

    # Gemini content parts (multimodal if images present)
    content_parts = [{"text": ANALYSIS_PROMPT + "\n\nCONTENT TO ANALYZE:\n" + combined_text}]
    for img in context["images"][:5]:   # Gemini supports up to 16 images
        content_parts.append({
            "inline_data": {"mime_type": img["mime"], "data": img["b64"]}
        })
        content_parts.append({"text": f"[Image above is: {img['name']}]"})

    response = generate_content_with_retry(
        client=client,
        model=MODEL,
        contents=[{"role": "user", "parts": content_parts}]
    )

    raw = response.text.strip()
    
    # Robust JSON parsing
    try:
        # Check for ```json ... ``` blocks
        match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", raw, re.DOTALL)
        if match:
            raw = match.group(1).strip()
        else:
            # Fallback to finding the first { and last }
            match_braces = re.search(r"(\{.*\})", raw, re.DOTALL)
            if match_braces:
                raw = match_braces.group(1).strip()
        return json.loads(raw)
    except Exception as e:
        print(f"⚠️ Failed parsing JSON response directly. Error: {e}")
        # Final fallback default structure
        return {
            "report_title": report_topic or "Data Analysis Report",
            "executive_summary": "An error occurred while parsing the Gemini analysis response.",
            "sections": [
                {
                    "title": "Raw Output Preview",
                    "content": raw[:2000].replace("\n", " "),
                    "key_insights": ["Could not parse JSON response from Gemini."]
                }
            ],
            "chart_specs": [],
            "conclusion": "Please verify your input files and try again.",
            "image_captions": {}
        }


# ── 3. Chart Generation ───────────────────────────────────────────────────────
def make_chart_matplotlib(spec: dict, dataframes: dict, temp_dir: Path, idx: int) -> str | None:
    """Create a beautiful Matplotlib chart and save as PNG. Returns file path string."""
    df = dataframes.get(spec.get("data_source"))
    if df is None and dataframes:
        df = list(dataframes.values())[0]   # fallback to first df
    if df is None:
        return None

    ctype   = spec.get("chart_type", "bar").lower()
    title   = spec.get("title", "Chart")
    colors  = ["#6366f1","#8b5cf6","#ec4899","#f59e0b","#10b981","#3b82f6","#ef4444"]

    fig, ax = plt.subplots(figsize=(8, 4.5), dpi=300)
    
    # Apply clean white design aesthetics
    sns.set_theme(style="white")
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#e2e8f0')
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.tick_params(colors='#475569', labelsize=9)
    ax.grid(axis='y', linestyle='--', alpha=0.5, color='#cbd5e1')

    try:
        if ctype == "pie":
            lc, vc = spec.get("label_col", df.columns[0]), spec.get("value_col", df.columns[1])
            if lc not in df.columns: lc = df.columns[0]
            if vc not in df.columns: vc = df.columns[1]
            
            grouped = df.groupby(lc)[vc].sum().reset_index()
            ax.pie(grouped[vc], labels=grouped[lc], autopct='%1.1f%%', colors=colors, 
                   startangle=140, wedgeprops={'edgecolor': 'w', 'linewidth': 1})
            ax.set_title(title, fontsize=13, fontweight='bold', pad=15, color='#1e293b')
            
        elif ctype == "bar":
            xc = spec.get("x_col", df.columns[0])
            yc = spec.get("y_col", df.columns[1])
            if xc not in df.columns: xc = df.columns[0]
            if yc not in df.columns: yc = df.columns[1]
            
            df_plot = df.copy()
            df_plot[yc] = pd.to_numeric(df_plot[yc], errors="coerce")
            df_plot = df_plot.dropna(subset=[yc])
            
            bar_colors = colors[:len(df_plot)] if len(df_plot) <= len(colors) else [colors[0]] * len(df_plot)
            ax.bar(df_plot[xc], df_plot[yc], color=bar_colors, edgecolor='none', width=0.55)
            ax.set_ylabel(yc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_xlabel(xc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_title(title, fontsize=13, fontweight='bold', pad=15, color='#1e293b')
            plt.xticks(rotation=30, ha='right')
            
        elif ctype == "line":
            xc = spec.get("x_col", df.columns[0])
            yc = spec.get("y_col", df.columns[1])
            if xc not in df.columns: xc = df.columns[0]
            if yc not in df.columns: yc = df.columns[1]
            
            df_plot = df.copy()
            df_plot[yc] = pd.to_numeric(df_plot[yc], errors="coerce")
            
            ax.plot(df_plot[xc], df_plot[yc], marker='o', linewidth=2.5, color=colors[0], markersize=6)
            ax.set_ylabel(yc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_xlabel(xc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_title(title, fontsize=13, fontweight='bold', pad=15, color='#1e293b')
            plt.xticks(rotation=30, ha='right')

        elif ctype == "scatter":
            xc = spec.get("x_col", df.columns[0])
            yc = spec.get("y_col", df.columns[1])
            if xc not in df.columns: xc = df.columns[0]
            if yc not in df.columns: yc = df.columns[1]
            
            df_plot = df.copy()
            df_plot[xc] = pd.to_numeric(df_plot[xc], errors="coerce")
            df_plot[yc] = pd.to_numeric(df_plot[yc], errors="coerce")
            
            ax.scatter(df_plot[xc], df_plot[yc], color=colors[2], s=70, alpha=0.8, edgecolors='white', linewidths=0.8)
            ax.set_xlabel(xc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_ylabel(yc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_title(title, fontsize=13, fontweight='bold', pad=15, color='#1e293b')

        elif ctype == "histogram":
            xc = spec.get("x_col", df.columns[0])
            if xc not in df.columns: xc = df.columns[0]
            
            vals = pd.to_numeric(df[xc], errors="coerce").dropna()
            ax.hist(vals, bins=15, color=colors[1], edgecolor='white', alpha=0.85)
            ax.set_xlabel(xc, fontsize=10, fontweight='bold', color='#475569')
            ax.set_ylabel("Frequency", fontsize=10, fontweight='bold', color='#475569')
            ax.set_title(title, fontsize=13, fontweight='bold', pad=15, color='#1e293b')

        elif ctype == "heatmap":
            numeric_df = df.select_dtypes(include="number")
            if not numeric_df.empty:
                sns.heatmap(numeric_df.corr(), annot=True, cmap="Purples", ax=ax, fmt=".2f", cbar=True)
                ax.set_title(title, fontsize=13, fontweight='bold', pad=15, color='#1e293b')
            else:
                plt.close(fig)
                return None
        else:
            plt.close(fig)
            return None

        plt.tight_layout()
        img_path = temp_dir / f"chart_{idx}.png"
        plt.savefig(img_path, format="png", dpi=300)
        plt.close(fig)
        return str(img_path)
    except Exception as e:
        print(f"  ⚠️  Matplotlib chart '{title}' failed: {e}")
        plt.close(fig)
        return None


# ── 4. DOCX Formatting Utilities ─────────────────────────────────────────────
def set_font(run, font_name="Arial", size_pt=11, color_rgb=(31, 41, 55), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level=1, primary_color=(49, 46, 129)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    size = 20 if level == 1 else (15 if level == 2 else 12)
    set_font(run, font_name="Arial", size_pt=size, color_rgb=primary_color, bold=True)
    return p

def add_callout(doc, title, text, bg_color="F3F4F6", border_color="4F46E5"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # Set cell margins (padding)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), '180') # Padding size in dxa
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    # Shading color
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    tcPr.append(shading)
    
    # Left border
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    # Add title and text
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(title)
    set_font(run_title, font_name="Arial", size_pt=12, color_rgb=(79, 70, 229), bold=True)
    
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    run_text = p2.add_run(text)
    set_font(run_text, font_name="Arial", size_pt=10, color_rgb=(55, 65, 81))

def add_table_from_df(doc, df):
    # Limit to first 12 columns/rows for safety in docx format
    df_trunc = df.iloc[:15, :10]
    
    table = doc.add_table(rows=1, cols=len(df_trunc.columns))
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row formatting
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df_trunc.columns):
        hdr_cells[i].text = str(col_name)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                set_font(run, font_name="Arial", size_pt=9.5, color_rgb=(255, 255, 255), bold=True)
        
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="312E81"/>') # Deep Indigo
        tcPr.append(shd)
        
    # Data rows formatting
    for r_idx, row in df_trunc.iterrows():
        row_cells = table.add_row().cells
        fill_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF" # light alternating colors
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                # Right-align numbers, left-align text
                try:
                    float(str(val).replace(",", "").replace("$", ""))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                except ValueError:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_font(run, font_name="Arial", size_pt=8.5, color_rgb=(31, 41, 55))
            
            tcPr = row_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
            tcPr.append(shd)

def build_docx(analysis: dict, chart_paths: list[str], context: dict, docx_path: Path, temp_dir: Path) -> None:
    doc = docx.Document()
    
    # Margins Setup
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(analysis.get("report_title", "Data Analysis Report"))
    set_font(run, font_name="Arial", size_pt=24, color_rgb=(49, 46, 129), bold=True)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    from datetime import date as dt
    date_str = dt.today().strftime("%B %d, %Y")
    run = subtitle_p.add_run(f"Prepared on {date_str} • Powered by Gemini AI")
    set_font(run, font_name="Arial", size_pt=10, color_rgb=(100, 116, 139), italic=True)
    
    # Executive Summary (Callout Block)
    exec_summary = analysis.get("executive_summary", "")
    if exec_summary:
        add_callout(doc, "EXECUTIVE SUMMARY", exec_summary)
        doc.add_paragraph() # spacing
        
    # Main Report Sections
    for sec in analysis.get("sections", []):
        add_heading(doc, sec.get("title", "Section"), level=1)
        
        # Split content into paragraphs
        content = sec.get("content", "")
        for para in content.split("\n\n"):
            if not para.strip(): continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(para.strip())
            set_font(run, font_name="Arial", size_pt=10.5, color_rgb=(31, 41, 55))
            
        # Insights list
        insights = sec.get("key_insights", [])
        if insights:
            for ins in insights:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(ins)
                set_font(run, font_name="Arial", size_pt=10, color_rgb=(31, 41, 55))
            doc.add_paragraph() # space after lists

    # Matplotlib Charts Section
    if chart_paths:
        add_heading(doc, "Data Visualizations", level=1)
        for i, (spec, path) in enumerate(zip(analysis.get("chart_specs", []), chart_paths)):
            if not path or not Path(path).exists(): continue
            
            add_heading(doc, spec.get("title", f"Visual Asset {i+1}"), level=2)
            
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.space_after = Pt(6)
            run = p_desc.add_run(spec.get("description", ""))
            set_font(run, font_name="Arial", size_pt=10, color_rgb=(100, 116, 139), italic=True)
            
            # Add chart image
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(12)
            p_img.paragraph_format.space_before = Pt(6)
            doc.add_picture(path, width=Inches(5.8))
            doc.add_paragraph()

    # User-uploaded Images Section
    captions = analysis.get("image_captions", {})
    if context["images"]:
        add_heading(doc, "Visual Attachments", level=1)
        for idx, img in enumerate(context["images"]):
            cap = captions.get(img["name"], img["name"])
            
            # Save raw b64 data to temporary file
            try:
                img_data = base64.b64decode(img["b64"])
                temp_img = temp_dir / f"user_img_{idx}{Path(img['name']).suffix}"
                temp_img.write_bytes(img_data)
                
                # Add picture
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(str(temp_img), width=Inches(5.0))
                
                # Add Caption below
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(12)
                run = p_cap.add_run(f"Figure: {cap}")
                set_font(run, font_name="Arial", size_pt=9, color_rgb=(100, 116, 139), italic=True)
            except Exception as e:
                print(f"  ⚠️  Failed to add user image {img['name']}: {e}")

    # Data Tables Section
    if context["dataframes"]:
        add_heading(doc, "Appendix: Detailed Data Tables", level=1)
        for name, df in context["dataframes"].items():
            add_heading(doc, f"Data Source: {name}", level=2)
            add_table_from_df(doc, df)
            doc.add_paragraph()

    # Conclusion & Recommendations
    conclusion_text = analysis.get("conclusion", "")
    if conclusion_text:
        add_heading(doc, "Conclusion & Recommendations", level=1)
        for para in conclusion_text.split("\n\n"):
            if not para.strip(): continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(para.strip())
            set_font(run, font_name="Arial", size_pt=10.5, color_rgb=(31, 41, 55))

    # Save DOCX
    doc.save(str(docx_path))
    print(f"✅  Word document saved → {docx_path}")


# ── 5. PDF Conversion via LibreOffice ────────────────────────────────────────
def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    """Run LibreOffice headless to convert the generated DOCX to PDF."""
    print("📄 Converting Word document to PDF using LibreOffice...")
    try:
        import subprocess
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            str(docx_path),
            "--outdir",
            str(output_dir)
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, check=True)
        pdf_path = output_dir / docx_path.with_suffix(".pdf").name
        if pdf_path.exists():
            print(f"✅  PDF successfully generated → {pdf_path}")
            return pdf_path
        else:
            print(f"⚠️  PDF not found at expected path: {pdf_path}")
            return None
    except Exception as e:
        print(f"⚠️  LibreOffice PDF conversion failed: {e}")
        return None


# ── Main ──────────────────────────────────────────────────────────────────────
def generate_report(file_paths: list[str], output_path_str: str, topic: str = "") -> None:
    out_path = Path(output_path_str)
    out_dir  = out_path.parent
    
    # Determine outputs based on extensions
    if out_path.suffix.lower() == ".pdf":
        docx_path = out_path.with_suffix(".docx")
        pdf_path  = out_path
        keep_docx = False
    elif out_path.suffix.lower() == ".docx":
        docx_path = out_path
        pdf_path  = out_path.with_suffix(".pdf")
        keep_docx = True
    else:
        # Default fallback
        docx_path = out_path.with_name(out_path.name + ".docx")
        pdf_path  = out_path.with_name(out_path.name + ".pdf")
        keep_docx = True

    with tempfile.TemporaryDirectory() as temp_dir_str:
        temp_dir = Path(temp_dir_str)
        
        print("🔍 Ingesting files...")
        context = ingest_files(file_paths)

        print("🤖 Analysing with Gemini...")
        analysis = analyze_with_gemini(context, topic)
        print(f"  📝 Title: {analysis.get('report_title')}")
        print(f"  📋 Sections: {len(analysis.get('sections', []))}")
        print(f"  📊 Charts planned: {len(analysis.get('chart_specs', []))}")

        print("📊 Generating charts using Matplotlib...")
        chart_paths = []
        for idx, spec in enumerate(analysis.get("chart_specs", [])):
            path = make_chart_matplotlib(spec, context["dataframes"], temp_dir, idx)
            if path:
                print(f"  ✅  {spec['chart_type'].upper()} chart generated: {spec['title']}")
                chart_paths.append(path)

        print("🎨 Assembling DOCX document...")
        build_docx(analysis, chart_paths, context, docx_path, temp_dir)
        
        # PDF Conversion
        pdf_res = convert_docx_to_pdf(docx_path, out_dir)
        
        # Clean up docx if user explicitly requested only a pdf
        if not keep_docx and docx_path.exists():
            docx_path.unlink()
            print(f"🗑️  Removed temporary DOCX file.")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python gemini_report_maker.py <output.docx|output.pdf> [topic] -- file1 file2 ...")
        sys.exit(1)

    out_file = sys.argv[1]
    topic    = ""
    files    = []
    rest     = sys.argv[2:]
    if rest and "--" in rest:
        sep   = rest.index("--")
        topic = " ".join(rest[:sep])
        files = rest[sep+1:]
    elif rest and not rest[0].startswith("--"):
        files = rest

    generate_report(files, out_file, topic)
